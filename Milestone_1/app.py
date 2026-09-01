import streamlit as st
import os
import ctypes
import sys
import tempfile
import shutil
import whisper
from moviepy import VideoFileClip
from google import genai
import chromadb
from sentence_transformers import SentenceTransformer
import soundfile as sf
import torch
from pyannote.audio import Pipeline
import io
import markdown
import docx
from xhtml2pdf import pisa

# --- WORKAROUNDS ---
try:
    pytorch_lib_dir = os.path.join(sys.prefix, "Lib", "site-packages", "torch", "lib")
    c10_dll = os.path.join(pytorch_lib_dir, "c10.dll")
    if os.path.exists(c10_dll):
        ctypes.CDLL(c10_dll)
except Exception:
    pass

try:
    import imageio_ffmpeg
    ffmpeg_src = imageio_ffmpeg.get_ffmpeg_exe()
    ffmpeg_dst = os.path.join(os.getcwd(), "ffmpeg.exe")
    if not os.path.exists(ffmpeg_dst):
        shutil.copy(ffmpeg_src, ffmpeg_dst)
except Exception:
    pass
# -------------------

st.set_page_config(
    page_title="SynthAI — Meeting Intelligence",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  DESIGN SYSTEM — injected CSS
# ─────────────────────────────────────────────
st.markdown("""
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">

<style>
/* ── ROOT VARIABLES ─────────────────────────────── */
:root {
  --bg-void:       #060812;
  --bg-base:       #0A0E1A;
  --bg-surface:    #0F1628;
  --bg-elevated:   #151c35;
  --glass-bg:      rgba(15, 22, 40, 0.65);
  --glass-border:  rgba(99, 102, 241, 0.18);
  --glass-border-hover: rgba(99, 102, 241, 0.45);

  --indigo:        #6366F1;
  --indigo-light:  #818CF8;
  --indigo-dark:   #4F46E5;
  --cyan:          #22D3EE;
  --cyan-light:    #67E8F9;
  --emerald:       #10B981;
  --amber:         #F59E0B;
  --rose:          #F43F5E;

  --text-primary:   #F0F4FF;
  --text-secondary: #94A3B8;
  --text-muted:     #4E5E7A;

  --radius-sm:  6px;
  --radius-md:  12px;
  --radius-lg:  20px;
  --radius-xl:  28px;

  --shadow-glow-indigo: 0 0 30px rgba(99, 102, 241, 0.25);
  --shadow-glow-cyan:   0 0 30px rgba(34, 211, 238, 0.20);
  --shadow-card:        0 4px 32px rgba(0,0,0,0.45);
  --transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1);
}

/* ── GLOBAL RESET ───────────────────────────────── */
*, *::before, *::after { box-sizing: border-box; }

html, body, [data-testid="stAppViewContainer"],
[data-testid="stApp"] {
  background: var(--bg-void) !important;
  font-family: 'Inter', sans-serif !important;
  color: var(--text-primary) !important;
}

[data-testid="stAppViewContainer"] {
  background:
    radial-gradient(ellipse 80% 60% at 20% -10%, rgba(99,102,241,0.12) 0%, transparent 60%),
    radial-gradient(ellipse 60% 40% at 80% 110%, rgba(34,211,238,0.08) 0%, transparent 55%),
    var(--bg-void) !important;
  min-height: 100vh;
}

[data-testid="stMain"] {
  background: transparent !important;
}

/* ── SIDEBAR ────────────────────────────────────── */
[data-testid="stSidebar"] {
  background: rgba(10, 14, 26, 0.92) !important;
  border-right: 1px solid var(--glass-border) !important;
  backdrop-filter: blur(20px) !important;
}

[data-testid="stSidebar"] > div {
  padding-top: 1.5rem !important;
}

/* sidebar title */
[data-testid="stSidebar"] h1 {
  font-size: 1rem !important;
  font-weight: 700 !important;
  letter-spacing: 0.05em !important;
  text-transform: uppercase !important;
  color: var(--indigo-light) !important;
  padding: 0 1rem 0.5rem !important;
  border-bottom: 1px solid var(--glass-border) !important;
  margin-bottom: 1.25rem !important;
}

/* sidebar labels */
[data-testid="stSidebar"] label {
  font-size: 0.72rem !important;
  font-weight: 600 !important;
  letter-spacing: 0.08em !important;
  text-transform: uppercase !important;
  color: var(--text-secondary) !important;
}

/* sidebar inputs */
[data-testid="stSidebar"] input[type="password"],
[data-testid="stSidebar"] input[type="text"] {
  background: rgba(99, 102, 241, 0.06) !important;
  border: 1px solid var(--glass-border) !important;
  border-radius: var(--radius-sm) !important;
  color: var(--text-primary) !important;
  font-family: 'JetBrains Mono', monospace !important;
  font-size: 0.8rem !important;
  transition: var(--transition) !important;
}

[data-testid="stSidebar"] input:focus {
  border-color: var(--indigo) !important;
  box-shadow: 0 0 0 3px rgba(99,102,241,0.15) !important;
  outline: none !important;
}

/* ── HEADINGS & TEXT ────────────────────────────── */
h1, h2, h3, h4 {
  font-family: 'Inter', sans-serif !important;
  color: var(--text-primary) !important;
  letter-spacing: -0.02em !important;
}

p, li, span, div {
  font-family: 'Inter', sans-serif !important;
}

/* ── HERO BANNER ────────────────────────────────── */
.hero-container {
  text-align: center;
  padding: 3.5rem 2rem 2rem;
  animation: fadeInDown 0.7s ease-out both;
}

.hero-logo {
  display: inline-flex;
  align-items: center;
  gap: 0.65rem;
  margin-bottom: 1rem;
}

.hero-logo-icon {
  width: 52px; height: 52px;
  background: linear-gradient(135deg, var(--indigo) 0%, var(--cyan) 100%);
  border-radius: 14px;
  display: flex; align-items: center; justify-content: center;
  font-size: 1.6rem;
  box-shadow: var(--shadow-glow-indigo);
}

.hero-brand {
  font-size: 2.4rem;
  font-weight: 800;
  background: linear-gradient(135deg, #F0F4FF 0%, var(--indigo-light) 45%, var(--cyan) 100%);
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
  background-clip: text;
  letter-spacing: -0.03em;
  line-height: 1;
}

.hero-tagline {
  font-size: 1.05rem;
  color: var(--text-secondary);
  font-weight: 400;
  margin-top: 0.5rem;
  letter-spacing: 0.01em;
}

.hero-badges {
  display: flex;
  gap: 0.5rem;
  justify-content: center;
  flex-wrap: wrap;
  margin-top: 1.25rem;
}

.hero-badge {
  display: inline-flex;
  align-items: center;
  gap: 0.35rem;
  padding: 0.3rem 0.75rem;
  border-radius: 100px;
  font-size: 0.72rem;
  font-weight: 600;
  letter-spacing: 0.04em;
  text-transform: uppercase;
  border: 1px solid;
}

.badge-whisper { background: rgba(99,102,241,0.1); border-color: rgba(99,102,241,0.3); color: var(--indigo-light); }
.badge-gemini  { background: rgba(34,211,238,0.1); border-color: rgba(34,211,238,0.3); color: var(--cyan-light); }
.badge-rag     { background: rgba(16,185,129,0.1); border-color: rgba(16,185,129,0.3); color: #6EE7B7; }

/* ── GLASS CARD ─────────────────────────────────── */
.glass-card {
  background: var(--glass-bg);
  border: 1px solid var(--glass-border);
  border-radius: var(--radius-lg);
  backdrop-filter: blur(20px);
  -webkit-backdrop-filter: blur(20px);
  padding: 1.75rem;
  box-shadow: var(--shadow-card);
  transition: var(--transition);
  animation: fadeInUp 0.5s ease-out both;
}

.glass-card:hover {
  border-color: var(--glass-border-hover);
  box-shadow: var(--shadow-glow-indigo), var(--shadow-card);
  transform: translateY(-2px);
}

/* ── UPLOAD ZONE ────────────────────────────────── */
.upload-zone-wrapper {
  background: linear-gradient(135deg, rgba(99,102,241,0.06) 0%, rgba(34,211,238,0.04) 100%);
  border: 2px dashed rgba(99,102,241,0.35);
  border-radius: var(--radius-xl);
  padding: 2.5rem 2rem;
  text-align: center;
  transition: var(--transition);
  position: relative;
  overflow: hidden;
  margin-bottom: 1rem;
}

.upload-zone-wrapper::before {
  content: '';
  position: absolute;
  inset: 0;
  background: linear-gradient(90deg, transparent, rgba(99,102,241,0.04), transparent);
  animation: shimmer 3s ease-in-out infinite;
}

.upload-icon {
  font-size: 3rem;
  margin-bottom: 0.75rem;
  display: block;
  animation: float 3s ease-in-out infinite;
}

.upload-title {
  font-size: 1.15rem;
  font-weight: 600;
  color: var(--text-primary);
  margin-bottom: 0.35rem;
}

.upload-subtitle {
  font-size: 0.82rem;
  color: var(--text-muted);
}

.format-chips {
  display: flex;
  gap: 0.4rem;
  justify-content: center;
  margin-top: 0.75rem;
  flex-wrap: wrap;
}

.format-chip {
  padding: 0.2rem 0.6rem;
  background: rgba(99,102,241,0.12);
  border: 1px solid rgba(99,102,241,0.25);
  border-radius: 100px;
  font-size: 0.7rem;
  font-weight: 600;
  color: var(--indigo-light);
  letter-spacing: 0.05em;
  font-family: 'JetBrains Mono', monospace;
}

/* ── STREAMLIT FILE UPLOADER OVERRIDE ───────────── */
[data-testid="stFileUploader"] {
  background: transparent !important;
}

[data-testid="stFileUploader"] > div {
  background: rgba(99, 102, 241, 0.05) !important;
  border: 2px dashed rgba(99, 102, 241, 0.35) !important;
  border-radius: var(--radius-lg) !important;
  transition: var(--transition) !important;
}

[data-testid="stFileUploader"] > div:hover {
  border-color: var(--indigo) !important;
  background: rgba(99, 102, 241, 0.10) !important;
}

[data-testid="stFileUploader"] label {
  color: var(--text-secondary) !important;
  font-size: 0.85rem !important;
}

[data-testid="stFileDropzone"] {
  background: transparent !important;
}

/* ── BUTTONS ────────────────────────────────────── */
[data-testid="stButton"] > button,
.stButton > button {
  background: linear-gradient(135deg, var(--indigo-dark) 0%, var(--indigo) 50%, #7C3AED 100%) !important;
  color: #fff !important;
  border: none !important;
  border-radius: var(--radius-md) !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 600 !important;
  font-size: 0.88rem !important;
  padding: 0.65rem 1.5rem !important;
  letter-spacing: 0.02em !important;
  cursor: pointer !important;
  transition: var(--transition) !important;
  box-shadow: 0 4px 20px rgba(99, 102, 241, 0.35) !important;
  position: relative !important;
  overflow: hidden !important;
}

[data-testid="stButton"] > button::after {
  content: '';
  position: absolute;
  inset: 0;
  background: linear-gradient(135deg, rgba(255,255,255,0.1), transparent);
  opacity: 0;
  transition: var(--transition);
}

[data-testid="stButton"] > button:hover {
  transform: translateY(-2px) !important;
  box-shadow: 0 8px 30px rgba(99, 102, 241, 0.55) !important;
}

[data-testid="stButton"] > button:hover::after {
  opacity: 1 !important;
}

[data-testid="stButton"] > button:active {
  transform: translateY(0) !important;
}

/* ── DOWNLOAD BUTTONS ───────────────────────────── */
[data-testid="stDownloadButton"] > button {
  background: rgba(15, 22, 40, 0.8) !important;
  color: var(--text-primary) !important;
  border: 1px solid var(--glass-border) !important;
  border-radius: var(--radius-md) !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 500 !important;
  font-size: 0.85rem !important;
  padding: 0.6rem 1.25rem !important;
  transition: var(--transition) !important;
  width: 100% !important;
}

[data-testid="stDownloadButton"] > button:hover {
  border-color: var(--indigo) !important;
  background: rgba(99, 102, 241, 0.12) !important;
  transform: translateY(-1px) !important;
  box-shadow: var(--shadow-glow-indigo) !important;
}

/* ── TABS ───────────────────────────────────────── */
[data-testid="stTabs"] {
  background: transparent !important;
}

[data-testid="stTabsList"] {
  background: rgba(10, 14, 26, 0.7) !important;
  border: 1px solid var(--glass-border) !important;
  border-radius: var(--radius-md) !important;
  padding: 4px !important;
  gap: 4px !important;
  backdrop-filter: blur(12px) !important;
  width: fit-content !important;
}

button[data-baseweb="tab"] {
  background: transparent !important;
  color: var(--text-muted) !important;
  border-radius: 8px !important;
  font-family: 'Inter', sans-serif !important;
  font-size: 0.85rem !important;
  font-weight: 500 !important;
  padding: 0.5rem 1.25rem !important;
  border: none !important;
  transition: var(--transition) !important;
  letter-spacing: 0.01em !important;
}

button[data-baseweb="tab"][aria-selected="true"] {
  background: linear-gradient(135deg, var(--indigo-dark), var(--indigo)) !important;
  color: #fff !important;
  box-shadow: 0 2px 12px rgba(99,102,241,0.4) !important;
}

button[data-baseweb="tab"]:hover:not([aria-selected="true"]) {
  color: var(--text-secondary) !important;
  background: rgba(99,102,241,0.08) !important;
}

[data-testid="stTabsContent"] {
  background: transparent !important;
  border: none !important;
  padding: 1.5rem 0 0 !important;
}

/* ── PROGRESS / STEPS ───────────────────────────── */
.steps-container {
  display: flex;
  justify-content: center;
  align-items: flex-start;
  gap: 0;
  margin: 1.5rem 0;
  padding: 1.5rem;
  background: var(--glass-bg);
  border: 1px solid var(--glass-border);
  border-radius: var(--radius-lg);
  backdrop-filter: blur(20px);
}

.step-item {
  display: flex;
  flex-direction: column;
  align-items: center;
  flex: 1;
  position: relative;
}

.step-item:not(:last-child)::after {
  content: '';
  position: absolute;
  top: 20px;
  left: 50%;
  width: 100%;
  height: 2px;
  background: var(--glass-border);
  z-index: 0;
}

.step-item.active:not(:last-child)::after {
  background: linear-gradient(90deg, var(--indigo), transparent);
}

.step-item.done:not(:last-child)::after {
  background: var(--indigo);
}

.step-dot {
  width: 40px; height: 40px;
  border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
  font-size: 1rem;
  position: relative; z-index: 1;
  transition: var(--transition);
}

.step-dot.waiting {
  background: rgba(78, 94, 122, 0.3);
  border: 2px solid var(--text-muted);
}

.step-dot.active {
  background: linear-gradient(135deg, var(--indigo-dark), var(--indigo));
  border: 2px solid var(--indigo-light);
  box-shadow: var(--shadow-glow-indigo);
  animation: pulse-glow 1.5s ease-in-out infinite;
}

.step-dot.done {
  background: linear-gradient(135deg, var(--emerald), #059669);
  border: 2px solid #6EE7B7;
}

.step-label {
  margin-top: 0.5rem;
  font-size: 0.7rem;
  font-weight: 600;
  letter-spacing: 0.05em;
  text-transform: uppercase;
  text-align: center;
  color: var(--text-muted);
}

.step-label.active { color: var(--indigo-light); }
.step-label.done   { color: var(--emerald); }

/* ── SECTION HEADERS ────────────────────────────── */
.section-header {
  display: flex;
  align-items: center;
  gap: 0.65rem;
  margin-bottom: 1.25rem;
}

.section-header-icon {
  width: 36px; height: 36px;
  border-radius: 10px;
  display: flex; align-items: center; justify-content: center;
  font-size: 1rem;
  flex-shrink: 0;
}

.icon-indigo { background: rgba(99,102,241,0.15); }
.icon-cyan   { background: rgba(34,211,238,0.12); }
.icon-emerald{ background: rgba(16,185,129,0.12); }
.icon-amber  { background: rgba(245,158,11,0.12); }

.section-header-text h3 {
  font-size: 1rem !important;
  font-weight: 700 !important;
  color: var(--text-primary) !important;
  margin: 0 !important;
}

.section-header-text span {
  font-size: 0.75rem;
  color: var(--text-muted);
  font-weight: 400;
}

/* ── REPORT CARDS ───────────────────────────────── */
.report-executive {
  background: linear-gradient(135deg, rgba(99,102,241,0.1) 0%, rgba(34,211,238,0.06) 100%);
  border: 1px solid rgba(99,102,241,0.25);
  border-radius: var(--radius-lg);
  padding: 1.75rem;
  margin-bottom: 1.25rem;
  position: relative;
  overflow: hidden;
}

.report-executive::before {
  content: '';
  position: absolute;
  top: 0; left: 0;
  width: 4px; height: 100%;
  background: linear-gradient(180deg, var(--indigo), var(--cyan));
}

.report-executive p, .report-executive ul, .report-executive li {
  color: var(--text-secondary) !important;
  line-height: 1.7 !important;
}

/* ── CHAT ───────────────────────────────────────── */
[data-testid="stChatMessage"] {
  background: transparent !important;
  border: none !important;
  padding: 0.5rem 0 !important;
}

[data-testid="stChatMessage"][data-testid*="user"] .stMarkdown,
[data-testid="stChatMessageContent"] {
  background: transparent !important;
}

.chat-message-user {
  background: linear-gradient(135deg, rgba(99,102,241,0.2), rgba(99,102,241,0.12)) !important;
  border: 1px solid rgba(99,102,241,0.25) !important;
  border-radius: 18px 18px 4px 18px !important;
  padding: 0.85rem 1.1rem !important;
  margin-left: auto !important;
  max-width: 80% !important;
  color: var(--text-primary) !important;
}

.chat-message-ai {
  background: rgba(15,22,40,0.7) !important;
  border: 1px solid var(--glass-border) !important;
  border-radius: 18px 18px 18px 4px !important;
  padding: 0.85rem 1.1rem !important;
  max-width: 90% !important;
  color: var(--text-primary) !important;
  backdrop-filter: blur(8px) !important;
}

/* Streamlit chat input */
[data-testid="stChatInput"] {
  background: rgba(15, 22, 40, 0.8) !important;
  border: 1px solid var(--glass-border) !important;
  border-radius: var(--radius-lg) !important;
  backdrop-filter: blur(12px) !important;
}

[data-testid="stChatInput"] textarea {
  background: transparent !important;
  color: var(--text-primary) !important;
  font-family: 'Inter', sans-serif !important;
  font-size: 0.9rem !important;
}

[data-testid="stChatInput"] textarea::placeholder {
  color: var(--text-muted) !important;
}

/* ── TRANSCRIPT CODE BLOCK ──────────────────────── */
[data-testid="stCode"],
.stCode {
  background: rgba(6, 8, 18, 0.8) !important;
  border: 1px solid var(--glass-border) !important;
  border-radius: var(--radius-md) !important;
}

pre, code {
  font-family: 'JetBrains Mono', monospace !important;
  font-size: 0.78rem !important;
  color: #CBD5E1 !important;
  background: transparent !important;
}

/* ── ALERTS / INFO BANNERS ──────────────────────── */
[data-testid="stAlert"] {
  border-radius: var(--radius-md) !important;
  border: 1px solid !important;
  backdrop-filter: blur(12px) !important;
}

[data-testid="stAlert"][data-baseweb="notification"] {
  background: rgba(15, 22, 40, 0.8) !important;
}

.stSuccess {
  background: rgba(16, 185, 129, 0.1) !important;
  border-color: rgba(16, 185, 129, 0.3) !important;
}

.stInfo {
  background: rgba(99, 102, 241, 0.1) !important;
  border-color: rgba(99, 102, 241, 0.3) !important;
}

.stWarning {
  background: rgba(245, 158, 11, 0.1) !important;
  border-color: rgba(245, 158, 11, 0.3) !important;
}

.stError {
  background: rgba(244, 63, 94, 0.1) !important;
  border-color: rgba(244, 63, 94, 0.3) !important;
}

/* ── SPINNER OVERRIDE ───────────────────────────── */
[data-testid="stSpinner"] > div {
  border-color: var(--indigo) transparent transparent transparent !important;
}

/* ── DIVIDERS ───────────────────────────────────── */
hr {
  border: none !important;
  border-top: 1px solid var(--glass-border) !important;
  margin: 1.5rem 0 !important;
}

/* ── SCROLLBAR ──────────────────────────────────── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb {
  background: rgba(99,102,241,0.3);
  border-radius: 100px;
}
::-webkit-scrollbar-thumb:hover { background: var(--indigo); }

/* ── METRIC CARDS ───────────────────────────────── */
.metric-card {
  background: var(--glass-bg);
  border: 1px solid var(--glass-border);
  border-radius: var(--radius-md);
  padding: 1.1rem 1.25rem;
  display: flex;
  align-items: center;
  gap: 0.85rem;
  transition: var(--transition);
}

.metric-card:hover {
  border-color: var(--glass-border-hover);
  transform: translateY(-1px);
}

.metric-icon {
  font-size: 1.5rem;
  width: 44px; height: 44px;
  display: flex; align-items: center; justify-content: center;
  border-radius: 10px;
  flex-shrink: 0;
}

.metric-label {
  font-size: 0.68rem;
  font-weight: 600;
  letter-spacing: 0.07em;
  text-transform: uppercase;
  color: var(--text-muted);
  margin-bottom: 0.15rem;
}

.metric-value {
  font-size: 1.15rem;
  font-weight: 700;
  color: var(--text-primary);
  letter-spacing: -0.01em;
}

/* ── STATUS DOT ─────────────────────────────────── */
.status-connected {
  display: inline-flex;
  align-items: center;
  gap: 0.4rem;
  padding: 0.3rem 0.7rem;
  background: rgba(16, 185, 129, 0.1);
  border: 1px solid rgba(16, 185, 129, 0.3);
  border-radius: 100px;
  font-size: 0.7rem;
  font-weight: 600;
  color: #6EE7B7;
  margin-top: 0.5rem;
}

.status-dot {
  width: 6px; height: 6px;
  border-radius: 50%;
  background: var(--emerald);
  animation: pulse-dot 2s ease-in-out infinite;
}

/* ── EXPORT SECTION ─────────────────────────────── */
.export-header {
  font-size: 0.72rem;
  font-weight: 700;
  letter-spacing: 0.09em;
  text-transform: uppercase;
  color: var(--text-muted);
  margin-bottom: 0.75rem;
  display: flex;
  align-items: center;
  gap: 0.4rem;
}

/* ── ANIMATIONS ─────────────────────────────────── */
@keyframes fadeInDown {
  from { opacity: 0; transform: translateY(-20px); }
  to   { opacity: 1; transform: translateY(0); }
}

@keyframes fadeInUp {
  from { opacity: 0; transform: translateY(20px); }
  to   { opacity: 1; transform: translateY(0); }
}

@keyframes float {
  0%, 100% { transform: translateY(0); }
  50%       { transform: translateY(-6px); }
}

@keyframes pulse-glow {
  0%, 100% { box-shadow: 0 0 0 0 rgba(99,102,241,0.5); }
  50%       { box-shadow: 0 0 0 8px rgba(99,102,241,0); }
}

@keyframes pulse-dot {
  0%, 100% { opacity: 1; }
  50%       { opacity: 0.4; }
}

@keyframes shimmer {
  0%   { transform: translateX(-100%); }
  50%  { transform: translateX(100%); }
  100% { transform: translateX(100%); }
}

@keyframes spin-glow {
  0%   { transform: rotate(0deg); }
  100% { transform: rotate(360deg); }
}

/* ── HIDE STREAMLIT CHROME ──────────────────────── */
#MainMenu, footer, header { visibility: hidden !important; }
[data-testid="stDecoration"] { display: none !important; }
[data-testid="stToolbar"]  { display: none !important; }

/* ── MARKDOWN INSIDE GLASS CARDS ────────────────── */
.element-container .stMarkdown p { color: var(--text-secondary) !important; line-height: 1.75 !important; }
.element-container .stMarkdown h2 { color: var(--text-primary) !important; margin-top: 1rem !important; }
.element-container .stMarkdown h3 { color: var(--indigo-light) !important; }
.element-container .stMarkdown li { color: var(--text-secondary) !important; }
.element-container .stMarkdown strong { color: var(--text-primary) !important; }

/* ── COLUMN GAPS ────────────────────────────────── */
[data-testid="stHorizontalBlock"] { gap: 1rem !important; }

/* ── PROCESS NEW FILE btn ───────────────────────── */
.new-file-btn > button {
  background: rgba(244, 63, 94, 0.08) !important;
  border: 1px solid rgba(244, 63, 94, 0.3) !important;
  color: #FDA4AF !important;
  box-shadow: none !important;
}

.new-file-btn > button:hover {
  background: rgba(244, 63, 94, 0.15) !important;
  border-color: rgba(244, 63, 94, 0.6) !important;
  transform: translateY(-1px) !important;
  box-shadow: 0 0 20px rgba(244, 63, 94, 0.2) !important;
}

/* selectbox */
[data-testid="stSelectbox"] select,
[data-baseweb="select"] {
  background: rgba(15, 22, 40, 0.8) !important;
  border-color: var(--glass-border) !important;
  color: var(--text-primary) !important;
}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="display:flex;align-items:center;gap:0.5rem;padding:0 1rem 1rem;">
      <div style="width:32px;height:32px;background:linear-gradient(135deg,#6366F1,#22D3EE);
                  border-radius:9px;display:flex;align-items:center;justify-content:center;
                  font-size:1.1rem;">🧠</div>
      <div style="font-size:1rem;font-weight:800;background:linear-gradient(135deg,#F0F4FF,#818CF8,#22D3EE);
                  -webkit-background-clip:text;-webkit-text-fill-color:transparent;
                  background-clip:text;letter-spacing:-0.02em;">SynthAI</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<p style="font-size:0.68rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;color:#4E5E7A;padding:0 0 0.5rem;">API Configuration</p>', unsafe_allow_html=True)

    gemini_key = st.text_input(
        "Gemini API Key",
        type="password",
        value="",
        help="Your Google Gemini API key for report generation and chat",
    )
    hf_token = st.text_input(
        "Hugging Face Token",
        type="password",
        value="",
        help="Required for Pyannote speaker diarization model",
    )

    keys_ok = bool(gemini_key and hf_token)

    if keys_ok:
        st.markdown("""
        <div class="status-connected">
          <div class="status-dot"></div>Keys Configured
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="display:inline-flex;align-items:center;gap:0.4rem;padding:0.3rem 0.7rem;
                    background:rgba(244,63,94,0.1);border:1px solid rgba(244,63,94,0.3);
                    border-radius:100px;font-size:0.7rem;font-weight:600;color:#FDA4AF;margin-top:0.5rem;">
          ⚠ Keys Required
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    st.markdown("""
    <p style="font-size:0.68rem;font-weight:700;letter-spacing:0.1em;text-transform:uppercase;
              color:#4E5E7A;margin-bottom:0.75rem;">Pipeline</p>
    <div style="display:flex;flex-direction:column;gap:0.5rem;">
      <div style="display:flex;align-items:center;gap:0.5rem;font-size:0.78rem;color:#94A3B8;">
        <span style="color:#6366F1">◆</span> Whisper — Transcription
      </div>
      <div style="display:flex;align-items:center;gap:0.5rem;font-size:0.78rem;color:#94A3B8;">
        <span style="color:#22D3EE">◆</span> Pyannote — Diarization
      </div>
      <div style="display:flex;align-items:center;gap:0.5rem;font-size:0.78rem;color:#94A3B8;">
        <span style="color:#10B981">◆</span> ChromaDB — RAG Index
      </div>
      <div style="display:flex;align-items:center;gap:0.5rem;font-size:0.78rem;color:#94A3B8;">
        <span style="color:#F59E0B">◆</span> Gemini — Intelligence
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    st.markdown("""
    <p style="font-size:0.65rem;color:#4E5E7A;line-height:1.6;text-align:center;">
      AI Meeting Synthesizer<br>
      <span style="color:#6366F1">v2.0</span> · Built with Streamlit
    </p>
    """, unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  GUARD
# ─────────────────────────────────────────────
if not keys_ok:
    st.markdown("""
    <div style="text-align:center;padding:4rem 2rem;">
      <div style="font-size:3rem;margin-bottom:1rem;">🔑</div>
      <h2 style="color:#F0F4FF;font-weight:700;margin-bottom:0.5rem;">API Keys Required</h2>
      <p style="color:#94A3B8;">Please enter your Gemini API Key and Hugging Face Token in the sidebar to continue.</p>
    </div>
    """, unsafe_allow_html=True)
    st.stop()


# ─────────────────────────────────────────────
#  MODEL CACHING
# ─────────────────────────────────────────────
@st.cache_resource
def load_whisper_model():
    return whisper.load_model("base")

@st.cache_resource
def load_diarization_pipeline(token):
    return Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", token=token)

@st.cache_resource
def load_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

@st.cache_resource
def get_chroma_client():
    return chromadb.Client()


# ─────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────
def format_timestamp(seconds):
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"[{mins:02d}:{secs:02d}]"

def merge_transcription_and_diarization(whisper_segments, diarization_result):
    merged_transcript = []
    for seg in whisper_segments:
        start, end, text = seg['start'], seg['end'], seg['text'].strip()
        speaker, max_overlap = "Unknown", 0
        for turn, _, spk in diarization_result.itertracks(yield_label=True):
            overlap = max(0, min(end, turn.end) - max(start, turn.start))
            if overlap > max_overlap:
                max_overlap, speaker = overlap, spk
        merged_transcript.append(f"{format_timestamp(start)} {speaker}: {text}")
    return "\n".join(merged_transcript)

def chunk_text(text, chunk_size=5):
    lines = text.split('\n')
    return ["\n".join(lines[i:i+chunk_size]) for i in range(0, len(lines), chunk_size)]

def generate_pdf_report(summary, transcript):
    html_content = f"""
    <html><head>
    <style>
      body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #1a1a2e; }}
      h1, h2, h3 {{ color: #4F46E5; }}
      .summary-block {{ background: #f0f0ff; border-left: 4px solid #6366F1;
                        padding: 1rem; border-radius: 4px; margin: 1rem 0; }}
      .transcript {{ font-family: monospace; white-space: pre-wrap; font-size: 11px;
                     background: #f8f8f8; padding: 1rem; border-radius: 4px; }}
      .meta {{ font-size: 11px; color: #666; border-top: 1px solid #eee; padding-top: 0.5rem; }}
    </style>
    </head><body>
      <h1>🧠 Meeting Intelligence Report</h1>
      <p class="meta">Generated by SynthAI Meeting Synthesizer</p>
      <div class="summary-block">{markdown.markdown(summary)}</div>
      <h2>Full Transcript</h2>
      <div class="transcript">{transcript}</div>
    </body></html>
    """
    pdf_buffer = io.BytesIO()
    pisa.CreatePDF(io.StringIO(html_content), dest=pdf_buffer)
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()

def generate_docx_report(summary, transcript):
    doc = docx.Document()
    doc.add_heading('Meeting Intelligence Report', 0)
    doc.add_heading('Summary', level=1)
    for line in summary.split('\n'):
        if not line.strip():
            continue
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith(('- ', '* ')):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    doc.add_heading('Full Transcript', level=1)
    for line in transcript.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()


# ─────────────────────────────────────────────
#  SESSION STATE
# ─────────────────────────────────────────────
for key, default in [
    ("transcript", None),
    ("chroma_collection", None),
    ("messages", []),
    ("report_summary", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default


# ─────────────────────────────────────────────
#  HERO SECTION
# ─────────────────────────────────────────────
st.markdown("""
<div class="hero-container">
  <div class="hero-logo">
    <div class="hero-logo-icon">🧠</div>
    <div class="hero-brand">SynthAI</div>
  </div>
  <div class="hero-tagline">Meeting Intelligence Platform — Transcribe, Analyze, Synthesize</div>
  <div class="hero-badges">
    <span class="hero-badge badge-whisper">⚡ Whisper Transcription</span>
    <span class="hero-badge badge-gemini">✦ Gemini Intelligence</span>
    <span class="hero-badge badge-rag">◈ RAG Chat</span>
  </div>
</div>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  UPLOAD SECTION
# ─────────────────────────────────────────────
if st.session_state.transcript is None:

    st.markdown("""
    <div style="max-width:680px;margin:0 auto 0.5rem;">
      <div class="upload-zone-wrapper">
        <span class="upload-icon">🎙️</span>
        <div class="upload-title">Drop your meeting recording here</div>
        <div class="upload-subtitle">Upload any audio or video file to begin AI-powered analysis</div>
        <div class="format-chips">
          <span class="format-chip">.mp3</span>
          <span class="format-chip">.wav</span>
          <span class="format-chip">.m4a</span>
          <span class="format-chip">.mp4</span>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    col_center = st.columns([1, 2, 1])[1]
    with col_center:
        uploaded_file = st.file_uploader(
            "Upload meeting recording",
            type=["mp3", "wav", "m4a", "mp4"],
            label_visibility="collapsed",
        )

    if uploaded_file is not None:
        # File info card
        file_size_mb = uploaded_file.size / (1024 * 1024)
        st.markdown(f"""
        <div style="max-width:680px;margin:0.75rem auto 1.25rem;
                    background:rgba(99,102,241,0.08);border:1px solid rgba(99,102,241,0.2);
                    border-radius:12px;padding:1rem 1.25rem;
                    display:flex;align-items:center;gap:0.75rem;">
          <span style="font-size:1.6rem;">🎵</span>
          <div style="flex:1;">
            <div style="font-weight:600;font-size:0.88rem;color:#F0F4FF;">{uploaded_file.name}</div>
            <div style="font-size:0.75rem;color:#94A3B8;margin-top:0.15rem;">{file_size_mb:.2f} MB · Ready for processing</div>
          </div>
          <div style="background:rgba(16,185,129,0.15);border:1px solid rgba(16,185,129,0.3);
                      border-radius:100px;padding:0.25rem 0.65rem;font-size:0.68rem;
                      font-weight:700;color:#6EE7B7;text-transform:uppercase;">Loaded</div>
        </div>
        """, unsafe_allow_html=True)

        col_btn = st.columns([1, 2, 1])[1]
        with col_btn:
            process_btn = st.button("🚀  Analyse Meeting", use_container_width=True)

        if process_btn:
            file_ext = os.path.splitext(uploaded_file.name)[1].lower()

            # ── STEP TRACKER ──────────────────────
            step_placeholder = st.empty()

            def render_steps(active_idx):
                steps = [
                    ("🎙", "Transcribing"),
                    ("👥", "Diarizing"),
                    ("🗃", "Indexing"),
                    ("✦", "Synthesizing"),
                ]
                items_html = ""
                for i, (icon, label) in enumerate(steps):
                    if i < active_idx:
                        dot_cls, lbl_cls = "done", "done"
                        icon_inner = "✓"
                    elif i == active_idx:
                        dot_cls, lbl_cls = "active", "active"
                        icon_inner = icon
                    else:
                        dot_cls, lbl_cls = "waiting", ""
                        icon_inner = icon
                    items_html += f"""
                    <div class="step-item {'done' if i < active_idx else ('active' if i == active_idx else '')}">
                      <div class="step-dot {dot_cls}">{icon_inner}</div>
                      <div class="step-label {lbl_cls}">{label}</div>
                    </div>
                    """
                step_placeholder.markdown(
                    f'<div class="steps-container">{items_html}</div>',
                    unsafe_allow_html=True,
                )

            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = tmp.name

                render_steps(0)
                whisper_model = load_whisper_model()
                diarization_pipeline = load_diarization_pipeline(hf_token)

                try:
                    audio_array = whisper.load_audio(tmp_path, sr=16000)
                    if len(audio_array) == 0:
                        st.error("The uploaded file does not contain audio (or the audio track is empty).")
                        st.stop()
                except Exception as e:
                    st.error(f"Could not extract audio: {e}")
                    st.stop()

                with st.spinner("Transcribing speech with Whisper…"):
                    result = whisper_model.transcribe(audio_array)

                render_steps(1)

                with st.spinner("Identifying speakers with Pyannote…"):
                    waveform_tensor = torch.from_numpy(audio_array).unsqueeze(0)
                    diarization = diarization_pipeline({"waveform": waveform_tensor, "sample_rate": 16000})

                render_steps(2)

                diarization_annotation = getattr(diarization, "speaker_diarization", diarization)
                final_transcript = merge_transcription_and_diarization(result['segments'], diarization_annotation)

                if not final_transcript:
                    st.error("Transcription returned empty output. Please try a different file.")
                    st.stop()

                st.session_state.transcript = final_transcript

                with st.spinner("Building RAG index in ChromaDB…"):
                    chroma_client = get_chroma_client()
                    try:
                        chroma_client.delete_collection("meeting_chunks")
                    except Exception:
                        pass
                    collection = chroma_client.create_collection("meeting_chunks")
                    embedder = load_embedding_model()
                    chunks = chunk_text(final_transcript)
                    embeddings = embedder.encode(chunks).tolist()
                    collection.add(
                        embeddings=embeddings,
                        documents=chunks,
                        ids=[f"chunk_{i}" for i in range(len(chunks))],
                    )
                    st.session_state.chroma_collection = collection

                render_steps(3)

                with st.spinner("Generating intelligence report with Gemini…"):
                    client = genai.Client(api_key=gemini_key)
                    prompt = f"""You are an expert AI meeting assistant. Analyze the following meeting transcript and generate a structured professional report.

The report MUST contain:
1. An **Executive Summary** (1-2 paragraphs)
2. **Key Discussion Points** (Bulleted list)
3. **Action Items & Owners** (Bulleted list of tasks and owners)

Meeting Transcript:
{final_transcript}
"""
                    import time
                    for attempt in range(3):
                        try:
                            summary_response = client.models.generate_content(
                                model='gemini-2.5-flash',
                                contents=prompt,
                            )
                            st.session_state.report_summary = summary_response.text
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 2:
                                time.sleep(2)
                            else:
                                st.session_state.report_summary = f"⚠ Failed to generate summary: {e}"

                render_steps(4)  # all done
                st.success("✅  Analysis complete! Scroll down to explore your report.")
                import time; time.sleep(0.8)
                st.rerun()

            except Exception as e:
                st.error(f"An unexpected error occurred: {e}")
            finally:
                if 'tmp_path' in locals() and os.path.exists(tmp_path):
                    try:
                        os.remove(tmp_path)
                    except Exception:
                        pass


# ─────────────────────────────────────────────
#  RESULTS VIEW
# ─────────────────────────────────────────────
if st.session_state.transcript:

    transcript_lines = st.session_state.transcript.strip().split('\n')
    speakers = set()
    for line in transcript_lines:
        parts = line.split(' ')
        if len(parts) > 1:
            spk = [p.rstrip(':') for p in parts if p.startswith('SPEAKER')]
            speakers.update(spk)

    # ── STATS ROW ──────────────────────────────
    st.markdown('<div style="height:0.5rem;"></div>', unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"""
        <div class="metric-card">
          <div class="metric-icon icon-indigo">🎙</div>
          <div>
            <div class="metric-label">Utterances</div>
            <div class="metric-value">{len(transcript_lines)}</div>
          </div>
        </div>""", unsafe_allow_html=True)
    with c2:
        word_count = len(st.session_state.transcript.split())
        st.markdown(f"""
        <div class="metric-card">
          <div class="metric-icon icon-cyan">💬</div>
          <div>
            <div class="metric-label">Words</div>
            <div class="metric-value">{word_count:,}</div>
          </div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""
        <div class="metric-card">
          <div class="metric-icon icon-emerald">👥</div>
          <div>
            <div class="metric-label">Speakers</div>
            <div class="metric-value">{max(len(speakers), 1)}</div>
          </div>
        </div>""", unsafe_allow_html=True)
    with c4:
        report_status = "Ready" if st.session_state.report_summary else "Pending"
        st.markdown(f"""
        <div class="metric-card">
          <div class="metric-icon icon-amber">✦</div>
          <div>
            <div class="metric-label">Report</div>
            <div class="metric-value">{report_status}</div>
          </div>
        </div>""", unsafe_allow_html=True)

    st.markdown('<div style="height:1rem;"></div>', unsafe_allow_html=True)

    # ── TABS ───────────────────────────────────
    tab1, tab2 = st.tabs(["📋  Meeting Intelligence", "💬  AI Chat Assistant"])

    # ════════════════════════════════════════════
    #  TAB 1 — REPORT
    # ════════════════════════════════════════════
    with tab1:

        if st.session_state.report_summary:

            # Executive Summary card
            st.markdown("""
            <div class="section-header">
              <div class="section-header-icon icon-indigo">✦</div>
              <div class="section-header-text">
                <h3>Meeting Intelligence Report</h3>
                <span>AI-generated analysis powered by Gemini</span>
              </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown('<div class="report-executive">', unsafe_allow_html=True)
            st.markdown(st.session_state.report_summary)
            st.markdown('</div>', unsafe_allow_html=True)

            st.markdown('<div style="height:0.5rem;"></div>', unsafe_allow_html=True)

            # Export section
            st.markdown("""
            <div class="export-header">
              ↓ &nbsp;Export Report
            </div>
            """, unsafe_allow_html=True)

            exp_col1, exp_col2, exp_col3 = st.columns(3)

            with exp_col1:
                pdf_bytes = generate_pdf_report(
                    st.session_state.report_summary, st.session_state.transcript
                )
                st.download_button(
                    label="📄  Download PDF",
                    data=pdf_bytes,
                    file_name="Meeting_Report.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )

            with exp_col2:
                docx_bytes = generate_docx_report(
                    st.session_state.report_summary, st.session_state.transcript
                )
                st.download_button(
                    label="📝  Download DOCX",
                    data=docx_bytes,
                    file_name="Meeting_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

            with exp_col3:
                st.download_button(
                    label="📃  Download Transcript",
                    data=st.session_state.transcript,
                    file_name="diarized_transcript.txt",
                    mime="text/plain",
                    use_container_width=True,
                )

        st.markdown("<hr>", unsafe_allow_html=True)

        # Full Transcript
        st.markdown("""
        <div class="section-header">
          <div class="section-header-icon icon-cyan">📜</div>
          <div class="section-header-text">
            <h3>Full Diarized Transcript</h3>
            <span>Speaker-attributed, timestamped</span>
          </div>
        </div>
        """, unsafe_allow_html=True)

        st.code(st.session_state.transcript, language="text")

        st.markdown("<hr>", unsafe_allow_html=True)

        # New file button
        st.markdown('<div class="new-file-btn">', unsafe_allow_html=True)
        if st.button("↺  Process a New File", use_container_width=False):
            for key in ["transcript", "chroma_collection", "report_summary", "messages"]:
                st.session_state[key] = None if key != "messages" else []
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    # ════════════════════════════════════════════
    #  TAB 2 — CHAT
    # ════════════════════════════════════════════
    with tab2:

        st.markdown("""
        <div class="section-header">
          <div class="section-header-icon icon-indigo">💬</div>
          <div class="section-header-text">
            <h3>AI Meeting Assistant</h3>
            <span>Ask anything about your meeting — powered by RAG + Gemini</span>
          </div>
        </div>
        """, unsafe_allow_html=True)

        # Chat history
        for msg in st.session_state.messages:
            role = msg["role"]
            with st.chat_message(role, avatar="🧑" if role == "user" else "🧠"):
                st.markdown(msg["content"])

        # Chat input
        if prompt := st.chat_input("Ask about the meeting… e.g. 'What were the key action items?'"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user", avatar="🧑"):
                st.markdown(prompt)

            with st.chat_message("assistant", avatar="🧠"):
                msg_placeholder = st.empty()

                embedder = load_embedding_model()
                query_embedding = embedder.encode([prompt]).tolist()

                results = st.session_state.chroma_collection.query(
                    query_embeddings=query_embedding,
                    n_results=3,
                )
                context = "\n\n".join(results['documents'][0])

                system_prompt = f"""You are an expert AI meeting assistant. Answer the user's question based ONLY on the provided meeting transcript context below.
If the information is not present in the context, say: "I don't have enough information from this meeting to answer that."
Be concise, professional, and precise.

Context:
{context}
"""
                try:
                    client = genai.Client(api_key=gemini_key)
                    import time
                    for attempt in range(3):
                        try:
                            response = client.models.generate_content_stream(
                                model='gemini-2.5-flash',
                                contents=system_prompt + f"\nUser Question: {prompt}",
                            )
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 2:
                                time.sleep(2)
                            else:
                                raise e

                    full_response = ""
                    for chunk in response:
                        if chunk.text:
                            full_response += chunk.text
                            msg_placeholder.markdown(full_response + " ▌")

                    msg_placeholder.markdown(full_response)
                    st.session_state.messages.append({"role": "assistant", "content": full_response})

                except Exception as e:
                    st.error(f"Could not generate response: {e}")
